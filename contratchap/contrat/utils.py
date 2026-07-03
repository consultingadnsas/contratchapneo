import docx
import re
import os
import subprocess
import shutil
from docxtpl import DocxTemplate

def extract_tags_grouped_by_paragraph(file_path):
    doc = docx.Document(file_path)
    pattern = re.compile(r'\{\{\s*(.*?)\s*\}\}')
    
    blocks_data = []
    
    def process_blocks(blocks):
        for block in blocks:
            text = block.text.strip()
            
            # Si le paragraphe est vide, on l'ajoute quand même pour conserver 
            # les sauts de ligne et la mise en page sur le frontend
            if not text:
                blocks_data.append({
                    "tags": [],
                    "context": ""
                })
                continue
            
            # 1. On cherche toutes les balises dans ce paragraphe
            matches = pattern.findall(text)
            
            # On nettoie les espaces autour des balises
            tags_in_this_block = [match.strip() for match in matches] if matches else []
            
            # 2. On ajoute TOUT le paragraphe à la liste, sans aucune troncature
            blocks_data.append({
                "tags": tags_in_this_block,
                "context": text
            })

    # Parcourir les paragraphes normaux du document
    process_blocks(doc.paragraphs)

    # Parcourir les tableaux (très fréquents dans les contrats)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_blocks(cell.paragraphs)

    return blocks_data

def convert_docx_to_pdf(docx_path, output_dir=None):
    """
    Convertit un fichier .docx en .pdf en utilisant LibreOffice en mode headless.
    Fonctionne sur Linux (Serveur, Docker), macOS et Windows.
    
    :param docx_path: Chemin absolu vers le fichier .docx à convertir
    :param output_dir: Dossier de destination du PDF (par défaut le même dossier que le docx)
    :return: Chemin absolu vers le fichier .pdf généré
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Le fichier source n'existe pas : {docx_path}")
    
    # Par défaut, on enregistre dans le même répertoire que le fichier docx
    if output_dir is None:
        output_dir = os.path.dirname(docx_path) or '.'

    # Détection automatique de l'exécutable LibreOffice (soffice)
    soffice_path = shutil.which('soffice') or shutil.which('libreoffice')
    
    # Sécurité / Fallbacks selon l'OS si non trouvé dans le PATH automatique
    if not soffice_path:
        if os.name == 'nt':  # Windows
            soffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
        elif os.path.exists('/Applications/LibreOffice.app/Contents/MacOS/soffice'):  # macOS
            soffice_path = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
        else:
            soffice_path = 'soffice'  # On tente la commande globale (Linux)

    # Commande magique de conversion headless de LibreOffice
    cmd = [
        soffice_path,
        '--headless',
        '--convert-to', 'pdf',
        '--outdir', output_dir,
        docx_path
    ]
    
    try:
        # Exécution de la commande système
        result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True, text=True)
        
        # Détermination du chemin du fichier PDF généré (LibreOffice garde le même nom de base)
        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        pdf_path = os.path.join(output_dir, f"{base_name}.pdf")
        
        if os.path.exists(pdf_path):
            return pdf_path
        else:
            raise Exception(f"Le fichier PDF n'a pas été généré. Sortie erreur : {result.stderr}")
            
    except subprocess.CalledProcessError as e:
        raise Exception(f"Erreur lors de la conversion LibreOffice : {e.stderr or e.stdout}")
    except FileNotFoundError:
        raise Exception(
            "LibreOffice n'est pas installé ou n'est pas accessible dans le PATH du système. "
            "Exécutez 'sudo apt-get install libreoffice' sur votre serveur Linux."
        )
    
def fill_docx_template(template_path, user_inputs, output_path):
    """
    Injecte les données de l'utilisateur dans le modèle DOCX.
    
    :param template_path: Chemin vers le fichier .docx original (le template avec les balises)
    :param user_inputs: Dictionnaire Python contenant les réponses (ex: {'nom_client': 'Lamine'})
    :param output_path: Chemin où sauvegarder le nouveau fichier .docx rempli
    :return: Le chemin vers le fichier rempli
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Le fichier modèle n'existe pas : {template_path}")

    # 1. On charge le template avec docxtpl
    doc = DocxTemplate(template_path)
    
    # 2. La magie opère : doc.render remplace toutes les balises {{ clé }} par les valeurs
    # Si une balise n'est pas dans user_inputs, elle sera simplement laissée vide ou effacée.
    doc.render(user_inputs)
    
    # 3. On sauvegarde le nouveau document prêt à être converti en PDF
    doc.save(output_path)
    
    return output_path